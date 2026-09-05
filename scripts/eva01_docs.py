# -*- coding: utf-8 -*-
"""Genera el alcance/EDT interno y el informe ejecutivo (máx. 7 planas)."""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from eva01_schedule import OUT as OUT_PLAN, run  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "evaluaciones" / "eva-01" / "informe"

MESES = [
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]
NAVY = RGBColor(0x10, 0x26, 0x3D)


def fecha(d: date) -> str:
    return f"{d.day} de {MESES[d.month - 1]} de {d.year}"


def set_run_font(run, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def setup_section(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    header = sec.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("Kiran · Plan preliminar del piloto · uso interno de dirección")
    set_run_font(r, size=9, color=RGBColor(0x55, 0x55, 0x55))
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Página ")
    set_run_font(r1, size=9, color=RGBColor(0x55, 0x55, 0x55))
    add_page_field(fp)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def p(doc, text, *, size=12, bold=False, italic=False, align="justify", space_after=6, first_line=True):
    para = doc.add_paragraph()
    if align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if first_line and align == "justify":
        para.paragraph_format.first_line_indent = Cm(0.75)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return para


def h(doc, text, level=1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    set_run_font(run, size=13 if level == 1 else 12, bold=True, color=NAVY)
    return para


def set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows, col_widths=None, header_fill="1F4E79"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(head)
        set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, header_fill)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = tbl.rows[r_i + 1].cells[c_i]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(str(val))
            set_run_font(run, size=9)
            if r_i % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    set_run_font(run, size=9, italic=True, color=RGBColor(0x44, 0x44, 0x44))


def ref(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Cm(1.25)
    para.paragraph_format.first_line_indent = Cm(-1.25)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run_font(run, size=12)


def by_code(tasks, code):
    return next(t for t in tasks if t.code == code)


def build_informe(info) -> Path:
    tasks = info["tasks"]
    end = info["end"]
    start = info["start"]
    days = info["finish_days"]
    t11 = by_code(tasks, "1.1")
    t23 = by_code(tasks, "2.3")
    t31 = by_code(tasks, "3.1")
    t44 = by_code(tasks, "4.4")
    t54 = by_code(tasks, "5.4")
    t73 = by_code(tasks, "7.3")
    m_inv = by_code(tasks, "8.1")
    m_tab = by_code(tasks, "8.2")
    m_tick = by_code(tasks, "8.3")
    m_rep = by_code(tasks, "8.4")
    m_fin = by_code(tasks, "8.5")

    doc = Document()
    setup_section(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    t.paragraph_format.line_spacing = 1.15
    r = t.add_run("Informe ejecutivo")
    set_run_font(r, size=16, bold=True, color=NAVY)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    t.paragraph_format.line_spacing = 1.15
    r = t.add_run("Kiran — plan preliminar del piloto de operación de kits solares")
    set_run_font(r, size=12, italic=True)

    meta = (
        "Producto: Kiran · “Visibilidad que mantiene la energía activa.”\n"
        "Organización: empresa social de energía solar comunitaria (India)\n"
        "Equipo de desarrollo: Giannina Guerrero · Nicolás Barra · Ari Araya · Skarlett Tropan"
    )
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(8)
    t.paragraph_format.line_spacing = 1.15
    r = t.add_run(meta)
    set_run_font(r, size=11)

    h(doc, "1. Contexto, necesidad de negocio y autorización")
    p(
        doc,
        "La empresa social enfrenta tres problemas de negocio: falta de energía confiable, "
        "desarrollo económico limitado y dependencia de combustibles fósiles. Su respuesta física "
        "son kits solares en una comunidad piloto. Kiran no es el kit: es el software que hace "
        "operable ese piloto. Construir Kiran es un proyecto (tiene inicio, fin y un resultado "
        "único). Cuando el técnico atienda tickets día a día, eso ya será operación. Si no se "
        "declara esa frontera, el alcance se confunde con electrificar la región.",
    )
    p(
        doc,
        "El caso de negocio tiene tres capas. Necesidad: kits instalados que se vuelven invisibles. "
        "Justificación: sin inventario vivo el subsidio se rinde con datos mentirosos. Estrategia: "
        "piloto de software híbrido, techo US$ 20.700. Output = Kiran. Outcome = fallas visibles. "
        "Valor = baja que saca el kit del activo, técnico que prioriza y donante que cree el informe "
        "(PMI, 2017; PMI, 2021).",
    )
    p(
        doc,
        "El acta de constitución autoriza Kiran —no la fábrica de paneles— y nombra a Giannina "
        f"Guerrero como directora del proyecto, con autoridad para usar recursos hasta el techo "
        f"de US$ 20.700. Con ella trabajan Nicolás Barra (backend), Ari Araya (infraestructura "
        f"y nube) y Skarlett Tropan (calidad e impacto). El proyecto inicia el {fecha(start)} y "
        f"cierra el {fecha(end)} ({days} días hábiles). Interesados: dirección de la empresa social "
        "(patrocinador), inversores y donantes, comunidad y hogares (no usan la app en esta fase), "
        "soporte técnico local y el partner tecnológico. Riesgos de alto nivel: conectividad "
        "irregular, datos incompletos de hogares y baja adopción del tablero en terreno.",
    )
    table(
        doc,
        ["Campo del acta", "Definición para este piloto"],
        [
            ["Propósito", "Construir Kiran: operar inventario, tablero, mantención y reporte de UNA comunidad piloto."],
            ["Objetivos medibles", "0 kits “fantasma”; tickets desde falla real; primer reporte con operación el " + fecha(t54.finish_date) + "."],
            ["Presupuesto resumido", "US$ 20.700 (contingencia 12 % incluida; costos hundidos fuera)."],
            ["Hitos resumidos", "Inventario " + fecha(m_inv.finish_date) + " · Tablero " + fecha(m_tab.finish_date) + " · Reporte " + fecha(m_rep.finish_date) + " · Cierre " + fecha(end) + "."],
            ["Criterio de aprobación", "Patrocinador acepta el primer reporte con datos reales; un kit dado de baja ya no cuenta como activo."],
        ],
        col_widths=[4.2, 11.8],
    )
    caption(doc, "Tabla 1. Extracto del acta de constitución (autoriza; no reemplaza el plan).")
    table(
        doc,
        ["Persona", "Rol en Kiran", "Qué sostiene en el desarrollo"],
        [
            ["Giannina Guerrero", "Directora del proyecto y frontend", "Acta, alcance, tablero operativo, relación con dirección"],
            ["Nicolás Barra", "Backend", "Modelo de datos, inventario, tickets, tratamiento de datos"],
            ["Ari Araya", "Infraestructura y nube", "Ambientes, despliegue, accesos técnicos, presupuesto de hosting"],
            ["Skarlett Tropan", "Calidad e impacto", "Pruebas, reportes a patrocinadores, capacitación al soporte local"],
        ],
        col_widths=[3.6, 4.6, 7.8],
    )
    caption(doc, "Tabla 2. Equipo de desarrollo del piloto.")

    h(doc, "2. Propuesta de valor y mejora sobre la solución del caso")
    p(
        doc,
        "La dirección ya pidió una plataforma de inventario, tablero, mantención y reportes. "
        "Copiar ese brief no basta: hay que mostrar cómo Kiran mejora esa solución y a quién le "
        "crea valor. El valor (PMI, 2021) es el indicador último de éxito: beneficios menos costos, "
        "visto por el interesado, y puede aparecer durante el piloto, al cierre o después.",
        first_line=True,
    )
    p(
        doc,
        "Propuesta de valor de Kiran: visibilidad que mantiene la energía activa. Para una "
        "empresa social que instala kits con plata de inversores y subsidio, Kiran es la capa "
        "de operación que vuelve observable, reparable y reportable lo ya instalado. A diferencia "
        "de una planilla o un grupo de WhatsApp, cada kit tiene estado vivo, cada falla nace como "
        "ticket asignable y cada informe a donantes se alimenta de operación real. Kiran no "
        "sustituye al panel: hace que el panel deje de ser un adorno en el techo.",
        first_line=True,
    )
    table(
        doc,
        ["Problemática de negocio", "Qué no resuelve el kit solo", "Cómo Kiran mejora la solución"],
        [
            [
                "Falta de energía confiable",
                "Un kit instalado y en silencio no ilumina: la falla no se ve ni se atiende",
                "Estados activo / falla / baja + ticket desde el kit: el técnico prioriza lo que de verdad está caído",
            ],
            [
                "Desarrollo económico limitado",
                "Sin energía estable no hay escuela ni taller; sin dato, la dirección no sabe dónde intervenir",
                "Tablero simple para soporte local: menos tiempo buscando el kit, más tiempo reparando",
            ],
            [
                "Dependencia de fósiles",
                "Si el kit “muerto” sigue en el inventario, el piloto miente el desplazamiento de diésel",
                "La baja saca el kit del activo. El reporte de impacto no cuenta fantasmas",
            ],
        ],
        col_widths=[3.6, 6.2, 6.2],
    )
    caption(doc, "Tabla 3. De la problemática de negocio a lo que Kiran mejora (sin fabricar paneles).")
    table(
        doc,
        ["Interesado", "Trabajo que necesita hacer", "Dolor hoy", "Valor que le entrega el piloto"],
        [
            [
                "Soporte local",
                "Encontrar y reparar kits",
                "Lista informal, WhatsApp, no sabe qué está en falla",
                "Tablero + ticket asignable; UX de terreno, no de oficina",
            ],
            [
                "Dirección / empresa social",
                "Operar el piloto y no mentir el activo",
                "Kits instalados “desaparecen” del radar",
                "Inventario vivo: activo, falla o baja, con dueño (hogar)",
            ],
            [
                "Inversores y donantes",
                "Ver evidencia de impacto",
                "Informes de intención, no de operación",
                "Reporte periódico alimentado por tickets y estados reales",
            ],
            [
                "Comunidad / hogares",
                "Tener luz que dure",
                "Esperan sin visibilidad de la falla",
                "Outcome: menos horas sin energía si la mantención se prioriza (indirecto)",
            ],
        ],
        col_widths=[3.2, 3.6, 4.4, 4.8],
    )
    caption(doc, "Tabla 4. Propuesta de valor por interesado.")
    p(
        doc,
        "Mejoras sobre el brief original, sin salir de inventario, tablero, tickets y reportes: "
        "cadena kit → estado → ticket → indicador; carga por lotes u offline; ticket como "
        "incidente de servicio; roles e ISO/IEC 27701; indicadores honestos (% activos, tiempo "
        "de resolución, bajas). No se agrega app ciudadana ni sensores masivos.",
        first_line=True,
    )

    h(doc, "3. Alcance, EDT, cronograma, recursos y costos")
    p(
        doc,
        "El enunciado de alcance se armó con producto, criterios de aceptación, entregables, "
        "exclusiones, restricciones y supuestos. El producto "
        "es Kiran, un sistema web (con carga por lotes u apoyo offline si falla la red) para "
        "registrar kits de una comunidad piloto, visualizar su estado, gestionar mantención "
        "y generar reportes periódicos. Criterios de aceptación: (1) cada kit tiene hogar "
        "dueño y estado activo / en falla / dado de baja, y la baja deja de contar como activo; "
        "(2) un ticket se abre desde un kit en falla y queda asignable al soporte local; "
        "(3) un reporte mensual exportable llega a patrocinadores con indicadores de operación "
        "real (no de maqueta). Esos criterios miden valor, no solo que “el módulo exista”.",
        first_line=True,
    )
    p(
        doc,
        "Entregables: módulos de inventario, tablero, tickets y reportes; documento de roles "
        "y privacidad; capacitación corta al soporte local; plan preliminar (este informe) y "
        "línea base en Project Libre. Restricciones: presupuesto del piloto, equipo de cuatro "
        "personas, conectividad de terreno y fecha del primer reporte a donantes. Supuestos: "
        "existe comunidad piloto identificada; hay al menos un técnico local; los inversores "
        "aceptan indicadores simples en esta fase; hay conectividad mínima o un procedimiento "
        "de carga por lotes. Si un supuesto cae, se trata como riesgo y no como sorpresa.",
    )
    p(
        doc,
        "Las exclusiones se cierran aquí para no hinchar el alcance: fabricación e instalación "
        "de paneles o kits físicos; tendido eléctrico; microfinanzas o cobro de tarifas; "
        "expansión a otras regiones; aplicación ciudadana masiva; operación continua posterior "
        "al piloto (eso es operación, no proyecto). Lo que no está en la EDT no se hace, o "
        "entra por control de cambios.",
    )
    p(
        doc,
        "Se generaron alternativas y se descartaron. App ciudadana masiva: cambia el usuario "
        "y el alcance. Tickets solo por WhatsApp: barato, pero no deja traza para ITIL ni para "
        "el reporte a donantes. Módulo interno de tickets + carga por lotes: es la opción que "
        "cabe en el acta. Una ronda tipo Delphi (cuestionario anónimo a técnico local, partner "
        "y dirección) sirve para no dejar que “el de más rango” imponga campos de más en el inventario.",
        first_line=True,
    )

    table(
        doc,
        ["EDT", "Cuenta de control", "Paquetes de trabajo"],
        [
            ["1", "Dirección del proyecto", "1.1 Acta e interesados · 1.2 Plan y línea base · 1.3 Riesgos y comunicación"],
            ["2", "Inventario de kits y hogares", "2.1 Modelo de datos · 2.2 Estados · 2.3 Carga inicial · 2.4 Validación local"],
            ["3", "Tablero operativo", "3.1 Vista general · 3.2 Vista individual · 3.3 Indicadores · 3.4 Ajustes de usabilidad"],
            ["4", "Monitoreo, tickets y mantención", "4.1 Rendimiento · 4.2 Flujo de tickets · 4.3 Asignación · 4.4 Prueba en terreno"],
            ["5", "Impacto y reportes", "5.1 Indicadores · 5.2 Informe periódico · 5.3 Exportación · 5.4 Primera entrega"],
            ["6", "Privacidad, accesos y despliegue", "6.1 Roles · 6.2 Datos de hogares · 6.3 Ambiente · 6.4 Puesta en marcha"],
            ["7", "Capacitación y transición", "7.1 Material · 7.2 Taller · 7.3 Traspaso operativo"],
        ],
        col_widths=[1.5, 5.2, 9.3],
    )
    caption(doc, "Tabla 5. EDT de primer nivel (el 100 % del trabajo acordado).")

    p(
        doc,
        "La EDT cubre el 100 % del trabajo del piloto. El diccionario traduce cada paquete "
        "a criterio, responsable, duración y costo. Extracto de dos paquetes de la ruta crítica:",
        first_line=True,
    )
    table(
        doc,
        ["Campo", "2.3 Carga inicial del piloto", "3.1 Vista general de la comunidad"],
        [
            ["Cuenta de control", "2 Inventario", "3 Tablero operativo"],
            ["Responsable", "Nicolás Barra (backend) + soporte local", "Giannina Guerrero (directora / frontend)"],
            [
                "Criterio de aceptación",
                "Los kits del piloto están cargados con hogar, comunidad y estado inicial verificable.",
                "El técnico ve, en una pantalla, recuento de kits por estado y acceso al detalle.",
            ],
            ["Supuesto", "El técnico entrega el listado maestro a tiempo.", "La carga 2.3 ya ocurrió; si no, el tablero es una demo vacía."],
            ["Duración", f"{t23.dur} días hábiles", f"{t31.dur} días hábiles"],
            ["Hito asociado", f"Inventario validado ({fecha(m_inv.finish_date)})", f"Tablero usable ({fecha(m_tab.finish_date)})"],
            ["Costo directo est.", "US$ 1.240 (carga + validación parcial)", "US$ 1.360 (frontend del tablero)"],
        ],
        col_widths=[3.4, 6.3, 6.3],
    )
    caption(doc, "Tabla 6. Extracto del diccionario de la EDT.")

    p(
        doc,
        "El cronograma se modeló en dos pases: primero duraciones brutas sin limitar recursos; "
        "después dependencias finish-to-start y calendario de lunes a viernes (8 h). El criterio "
        "técnico es el método de la ruta crítica (CPM): el camino más largo determina la duración "
        "mínima. Ese camino es 1.1–2.1–2.2–2.3–3.1–3.2–4.2–4.3–4.4–5.4–7.3. En esta operación: "
        "sin modelo ni carga de inventario no hay tablero con datos reales; sin tablero individual "
        "no hay flujo de tickets útil; sin prueba en terreno el reporte a donantes es una infografía "
        f"vacía. Duración del proyecto: {days} días hábiles ({fecha(start)} al {fecha(end)}). "
        "Cualquier atraso en esas actividades (holgura 0) mueve el fin. El camino más largo "
        "es el que manda, no el que “se siente más importante”.",
        first_line=True,
    )
    table(
        doc,
        ["Camino", "Secuencia", "Suma (días)", "¿Crítica?"],
        [
            [
                "A Inventario–tablero–tickets",
                "1.1–2.1–2.2–2.3–3.1–3.2–4.2–4.3–4.4–5.4–7.3",
                str(days),
                "Sí (holgura 0)",
            ],
            [
                "B Privacidad y despliegue",
                "1.1–6.1–6.2–6.3–6.4 (4.4 igual espera a 4.3)",
                "23",
                "No (holgura alta)",
            ],
            [
                "C Indicadores de impacto",
                "1.1–5.1–5.2–5.3 (5.4 espera la prueba 4.4)",
                "17",
                "No (reporte sin operación real no vale)",
            ],
        ],
        col_widths=[4.4, 6.6, 2.6, 2.4],
    )
    caption(doc, "Tabla 7. Caminos analizados con CPM. El camino A mueve la fecha de entrega.")
    p(
        doc,
        "Donde hay incertidumbre de terreno se usó PERT. Para la prueba en terreno (4.4): "
        "optimista 5 días, más probable 7, pesimista 13. Duración esperada = (5 + 4×7 + 13) / 6 = "
        "8 días; desviación = (13 − 5) / 6 ≈ 1,3 días. A gerencia se muestran hitos; al equipo, "
        "el Gantt de Project Libre.",
        first_line=True,
    )
    table(
        doc,
        ["Hito", "Fecha", "Para quién"],
        [
            [f"Inventario validado", fecha(m_inv.finish_date), "Equipo y soporte local"],
            [f"Tablero usable", fecha(m_tab.finish_date), "Técnico local y dirección"],
            [f"Tickets en terreno", fecha(m_tick.finish_date), "Soporte local"],
            [f"Primer reporte al patrocinador", fecha(m_rep.finish_date), "Inversores y donantes"],
            [f"Piloto entregado", fecha(m_fin.finish_date), "Dirección / cierre de fase"],
        ],
        col_widths=[6.5, 4.5, 5.0],
    )
    caption(doc, "Tabla 8. Cronograma de hitos (para dirección).")

    p(
        doc,
        "Los recursos se desglosaron en una RBS. Personas: Giannina Guerrero, directora del "
        "proyecto y frontend (0,5 FTE en dirección más el tablero); Nicolás Barra, backend; "
        "Ari Araya, infraestructura y nube; Skarlett Tropan, calidad e impacto (0,6 FTE); y "
        "soporte local (0,4 FTE). Tecnología: nube, repositorio, Project Libre, tablero Kanban "
        "y un punto de carga en la comunidad. Financieros: aporte de inversores de impacto, "
        "subvención y reserva de contingencia. La disponibilidad del técnico local es el recurso "
        "más frágil: si no carga datos, la ruta crítica se rompe aunque el código esté listo.",
        first_line=True,
    )
    table(
        doc,
        ["Tipo de costo", "Ejemplo en el piloto", "Monto (US$)"],
        [
            ["Fijo", "Remuneración de la directora de proyecto (0,5 FTE × 3 meses)", "2.400"],
            ["Variable", "Horas de soporte local según kits cargados; hosting según uso", "1.240"],
            ["Directo", "Equipo de desarrollo, nube del piloto, taller de capacitación", "14.470"],
            ["Indirecto", "Administración / PMO de la empresa social (cuota del piloto)", "1.380"],
            ["De oportunidad", "El mismo equipo no construye un módulo de cobro (excluido a propósito)", "No caja"],
            ["Hundido", "Estudio de terreno previo ya pagado; no decide continuar o parar", "2.800 (fuera)"],
            ["Contingencia (12 %)", "Reserva sobre costos directos + variables del piloto", "1.890"],
            ["Presupuesto autorizado", "Techo del acta (hundidos excluidos)", "20.700"],
        ],
        col_widths=[3.6, 10.0, 2.4],
    )
    caption(doc, "Tabla 9. Presupuesto inicial por tipo de costo. Los hundidos no entran a la decisión.")

    h(doc, "4. Factores ambientales, estándares e impacto en el caso")
    p(
        doc,
        "Los factores ambientales de la empresa (EEF) son condiciones que el equipo no controla "
        "y que entran a la planificación (Project Management Institute [PMI], 2017). Internos: "
        "cultura de transparencia exigida por donantes; estructura matricial débil (empresa social "
        "+ partner tecnológico + un equipo de software de cuatro personas); infraestructura de nube aún no institucionalizada; "
        "capacidad limitada del equipo. Externos: financiamiento mixto; distancia cultural y geográfica "
        "con la comunidad piloto; marco legal de datos de hogares; conectividad física irregular; "
        "estándares de la industria. Los OPA (plantillas de acta y EDT, este plan, lecciones del "
        "equipo en Trello) sí se pueden usar; el clima, no. Los interesados pueden mover alcance, plazo, costo, equipo, calidad y la "
        "definición de éxito (PMI, 2021). Un donante que pide el reporte antes comprime 5.4; un "
        "técnico que no carga datos rompe 2.3 y, con ello, toda la ruta crítica. Por eso los "
        "canales se declaran ahora: hitos para inversores, taller para el técnico, roles para "
        "privacidad. No se involucra al interesado “al cierre”.",
        first_line=True,
    )
    table(
        doc,
        ["Factor o estándar", "Tipo", "Impacto concreto en el piloto"],
        [
            ["Financiamiento mixto (inversores + subvención)", "EEF externo financiero", "Hitos de reporte 8.4 inamovibles: el cronograma se diseña hacia esa fecha, no al revés."],
            ["Cultura comunitaria e idioma", "EEF externo social", "UX simple, pocos campos, capacitación 7.x; no se asume onboarding “tipo app urbana”."],
            ["Conectividad irregular", "EEF físico", "Supuesto de carga por lotes; si cae, el inventario miente y la ruta crítica se rompe."],
            ["Datos de hogares", "EEF legal", "Minimizar datos, roles 6.1 y tratamiento 6.2; alineación con ISO/IEC 27701."],
            ["PMBOK 6.ª", "Estándar de dirección", "Acta, EDT, CPM, interesados y línea base. Dirige el proyecto, no la operación diaria."],
            ["CMMI", "Madurez de procesos", "El piloto apunta a disciplina básica (plan, evidencias), no a un nivel 4 o 5."],
            ["ITIL 4", "Gestión de servicio", "El módulo de tickets se diseña como incidente/mantención, no como formulario suelto."],
            ["COBIT", "Gobierno de TI", "Quién aprueba accesos y reportes hacia inversores; evita que “el de sistemas” decida solo."],
        ],
        col_widths=[4.4, 3.2, 8.4],
    )
    caption(doc, "Tabla 10. EEF y estándares con impacto concreto en este piloto.")

    h(doc, "5. Estrategia de planificación e integración a la organización")
    p(
        doc,
        "Se descarta un enfoque 100 % predictivo: el flujo real del técnico local no está cerrado. "
        "Se descarta un enfoque 100 % adaptativo: los estados del kit, la privacidad y el calendario "
        "de donantes no pueden “descubrirse” cada semana. La estrategia es híbrida, calzada al terreno "
        "de esta organización (Wysocki, 2019; PMI, 2017).",
        first_line=True,
    )
    p(
        doc,
        "Capa predictiva (receta): inventario y estados del kit (cuenta 2), roles y privacidad "
        "(cuenta 6) y el calendario de reportes a donantes (hitos de la cuenta 5). Capa adaptativa "
        "(GPS): UX del tablero (cuenta 3), flujo de tickets (cuenta 4), cómo se miden los indicadores "
        "de impacto y el trabajo con soporte local (cuenta 7). La dirección (cuenta 1) integra: una "
        "línea base de hitos y un backlog semanal en Trello debajo. Eso es hibridación según el "
        "terreno de esta organización, no moda del equipo (PMI, 2017; Wysocki, 2019).",
    )
    p(
        doc,
        "Integración organizacional: el piloto no vive en el vacío del equipo de software. Se "
        "inserta en una empresa social que debe ser ágil en terreno y estricta con plata ajena. "
        "Por eso el plan incorpora canales de comunicación a inversores (hitos), roles de acceso "
        "(COBIT + privacidad), y una transición al soporte local para no dejar Kiran huérfano "
        "al cerrar. Un sprint de dos semanas (SCRUMstudy, 2023) solo para tablero y tickets; si "
        "aparece un campo nuevo de inventario, es cambio a la línea base, no un sticker en Trello.",
    )
    table(
        doc,
        ["Cuenta", "Giannina (DP/front)", "Nicolás (backend)", "Ari (infra)", "Skarlett (QA/impacto)"],
        [
            ["1 Dirección", "R/A", "C", "C", "C"],
            ["2 Inventario", "A", "R", "C", "C"],
            ["3 Tablero", "R/A", "C", "I", "C"],
            ["4 Tickets", "A", "R", "C", "C"],
            ["5 Impacto", "A", "C", "I", "R"],
            ["6 Privacidad/despliegue", "A", "C", "R", "C"],
            ["7 Capacitación", "A", "C", "C", "R"],
        ],
        col_widths=[3.6, 3.1, 3.1, 3.2, 3.0],
    )
    caption(doc, "Tabla 11. RACI del equipo. R = hace · A = rinde · C = consulta · I = informa.")

    h(doc, "6. Herramientas utilizadas y justificación")
    p(
        doc,
        "La herramienta sigue a la estrategia y a lo que esta organización puede usar de verdad, "
        "no a la moda del equipo. Para la línea base predictiva se eligió Project Libre: cubre EDT, "
        "precedencias, ruta crítica, recursos y Gantt, sin licencia de Microsoft Project, coherente "
        "con un piloto de impacto y presupuesto acotado. Office 365 (Word y Excel) sostiene el "
        "diccionario de la EDT, el presupuesto por tipo de costo y este informe ejecutivo. Un tablero "
        "Kanban (Trello) cubre la capa adaptativa de tablero y tickets. La pizarra se usa para "
        "alinear a dirección y al técnico local; no es la línea base.",
        first_line=True,
    )
    table(
        doc,
        ["Herramienta", "Dónde se usó en este plan", "Por qué esta organización"],
        [
            ["Project Libre", "Línea base: EDT, CPM, recursos, hitos (Cronograma-Kiran.xml)", "Hay que defender plazo de subsidio; no hay licencia MS Project."],
            ["Excel / Office 365", "Diccionario EDT, RBS y presupuesto (Tabla 9)", "Formato que dirección y donantes ya leen; no reemplaza la ruta crítica."],
            ["Trello (Kanban)", "Capa adaptativa de cuentas 3 y 4 (tablero y tickets)", "Equipo de cuatro; el flujo se descubre en terreno y no cabe congelarlo en el Gantt."],
            ["Word (este informe)", "Plan que dirección lee para autorizar el piloto", "Formato que ya usan patrocinadores; las citas marcan de dónde sale cada criterio."],
            ["Pizarra", "Alineación con dirección y técnico local", "Entendimiento común en una reunión; no sustituye Project Libre."],
        ],
        col_widths=[3.6, 6.4, 6.0],
    )
    caption(doc, "Tabla 12. Selección de herramientas según lo que esta organización puede usar de verdad.")

    p(
        doc,
        "En síntesis: el piloto queda justificado por valor (mejora la solución energética de "
        "esta organización sin fabricar paneles), autorizado (acta), limitado (exclusiones), partido (EDT), "
        "fechado por el camino más largo (CPM + PERT), costoso con tipos y contingencia, honesto "
        "con el clima (EEF) y caminado en híbrido. El éxito no será “haber programado”, sino "
        "kits observables, fallas atendibles y un reporte que un donante pueda creer.",
        first_line=True,
    )

    h(doc, "Referencias")
    ref(
        doc,
        "Baud, J.-L. (2020). ITIL 4: Entender el enfoque y adoptar las buenas prácticas. Ediciones ENI.",
    )
    ref(
        doc,
        "Project Management Institute. (2017). Guía de los fundamentos para la dirección de proyectos (Guía del PMBOK) (6.ª ed.).",
    )
    ref(
        doc,
        "Project Management Institute. (2021). Guía de los fundamentos para la dirección de proyectos y El estándar para la dirección de proyectos (7.ª ed.).",
    )
    ref(
        doc,
        "SCRUMstudy. (2023). Guía SBOK (4.ª ed.).",
    )
    ref(
        doc,
        "Wysocki, R. K. (2019). Effective project management: Traditional, agile, extreme, hybrid (8th ed.). Wiley.",
    )

    path = OUT / "Informe-Ejecutivo-Eva01.docx"
    try:
        doc.save(path)
    except PermissionError:
        path = OUT / "Informe-Ejecutivo-Eva01-actualizado.docx"
        doc.save(path)
    return path


def build_alcance(info) -> Path:
    tasks = info["tasks"]
    doc = Document()
    setup_section(doc)
    # override header
    doc.sections[0].header.paragraphs[0].clear()
    r = doc.sections[0].header.paragraphs[0].add_run(
        "Kiran · Alcance cerrado, exclusiones y EDT · uso interno de dirección"
    )
    set_run_font(r, size=9, color=RGBColor(0x55, 0x55, 0x55))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Kiran — alcance cerrado, exclusiones y EDT")
    set_run_font(r, size=16, bold=True, color=NAVY)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(8)
    r = t.add_run("Visibilidad que mantiene la energía activa.")
    set_run_font(r, size=12, italic=True)

    h(doc, "Qué es este proyecto (y qué no)")
    p(
        doc,
        "Este documento cierra el alcance ANTES del Gantt. Si el árbol y las exclusiones no "
        "están firmados por el equipo, cualquier barra del cronograma es un deseo. Kiran es el "
        "proyecto de software de operación del piloto. No es un proyecto de ingeniería eléctrica.",
    )
    h(doc, "Propuesta de valor (por qué este software, y no solo “otra planilla”)")
    p(
        doc,
        "Esta organización enfrenta falta de energía confiable, desarrollo económico limitado y "
        "dependencia de fósiles. Los kits solares son la solución física. Kiran no compite con "
        "el panel: mejora esa solución para que un kit instalado no se vuelva invisible. "
        "Valor = kits observables + fallas atendibles + reportes creíbles. Promesa: visibilidad "
        "que mantiene la energía activa. Si el técnico no usa el tablero, no hay valor aunque "
        "el código compile.",
    )
    p(
        doc,
        "Mejoras sobre el brief de dirección, sin hinchar alcance: cadena kit→ticket→indicador; "
        "carga por lotes/offline; ticket como incidente de servicio; privacidad de hogares; "
        "indicadores honestos (activos, tiempo de resolución, bajas). Fuera: app ciudadana, "
        "sensores masivos, paneles.",
        first_line=True,
    )

    h(doc, "Equipo de desarrollo")
    table(
        doc,
        ["Persona", "Rol", "Sostiene"],
        [
            ["Giannina Guerrero", "Directora del proyecto y frontend", "Acta, alcance, tablero, relación con dirección"],
            ["Nicolás Barra", "Backend", "Datos, inventario, tickets, tratamiento de hogares"],
            ["Ari Araya", "Infraestructura y nube", "Ambientes, despliegue, accesos técnicos, hosting"],
            ["Skarlett Tropan", "Calidad e impacto", "Pruebas, reportes a patrocinadores, capacitación"],
        ],
        col_widths=[3.6, 4.8, 7.6],
    )

    h(doc, "Enunciado de alcance")
    table(
        doc,
        ["Elemento", "Definición cerrada para el piloto"],
        [
            [
                "Producto",
                "Kiran: sistema web para registrar kits de UNA comunidad piloto, ver estado, gestionar mantención y emitir reportes. Apoyo de carga por lotes u offline si la red falla.",
            ],
            [
                "Criterios de aceptación",
                "Kit con hogar + estado (activo/falla/baja). Ticket nace de un kit en falla y se asigna. Reporte mensual exportable a patrocinadores.",
            ],
            [
                "Entregables",
                "Módulos 2–5, roles/privacidad (6), capacitación (7), acta y plan preliminar, línea base en Project Libre.",
            ],
            [
                "Exclusiones",
                "Ver tabla siguiente. Quedan fuera de la EDT y de la ruta crítica.",
            ],
            [
                "Restricciones",
                "Techo US$ 20.700; equipo de 4 + técnico local; conectividad irregular; primer reporte a donantes el "
                + fecha(by_code(tasks, "8.4").finish_date)
                + ".",
            ],
            [
                "Supuestos",
                "Comunidad piloto identificada; hay un técnico local; inversores aceptan indicadores simples; existe procedimiento de carga si no hay señal.",
            ],
        ],
        col_widths=[4.0, 12.0],
    )

    h(doc, "Exclusiones cerradas (con el porqué)")
    table(
        doc,
        ["Queda fuera", "Por qué se excluye", "Qué pasaría si no se escribe"],
        [
            ["Fabricar o instalar paneles y kits", "Es infraestructura energética, no software", "El alcance se hincha y el CPM deja de ser de este producto"],
            ["Tendido eléctrico / microred física", "Fuera de la competencia de este equipo de desarrollo", "Se promete algo que no se puede entregar en el plazo del piloto"],
            ["Microfinanzas o cobro de tarifas", "No está en las funcionalidades pedidas por dirección", "Nace un segundo producto (pagos) sin presupuesto"],
            ["Expansión a otras regiones", "El acta autoriza UNA comunidad piloto", "Se planifica un roll-out que nadie financió"],
            ["App ciudadana masiva", "El usuario de esta fase es el soporte local y la dirección", "UX y alcance cambian de destinatario"],
            ["Operación continua post-piloto", "Eso es operación, no proyecto", "El cierre del piloto no tendría frontera"],
        ],
        col_widths=[4.4, 5.8, 5.8],
    )
    p(
        doc,
        "Regla del equipo: si alguien pide “de paso un dashboard para toda India”, es un cambio. "
        "No se mete por el ladito al Gantt.",
        first_line=True,
    )

    h(doc, "EDT completa (regla del 100 %)")
    p(
        doc,
        "Nivel 1 = cuentas de control (ahí se medirá alcance, plazo y costo juntos). "
        "Nivel 2 = paquetes de trabajo estimables. Lo que no aparece aquí no existe para el piloto.",
        first_line=True,
    )
    rows = []
    for t in tasks:
        if t.code.startswith("8"):
            continue
        if t.summary:
            rows.append([t.code, t.name.upper(), "Cuenta de control", "—", "—"])
        else:
            crit = "Sí" if t.slack == 0 else "No"
            rows.append(
                [
                    t.code,
                    t.name,
                    "Paquete de trabajo",
                    f"{t.dur} d",
                    crit,
                ]
            )
    table(
        doc,
        ["ID", "Nombre", "Nivel", "Duración", "¿Crítica?"],
        rows,
        col_widths=[2.0, 8.2, 3.4, 2.0, 2.4],
    )
    caption(doc, "Tabla. EDT numerada. Las actividades críticas alimentan el XML de Project Libre.")

    h(doc, "Dependencia que amarra el árbol al calendario")
    p(
        doc,
        "Inventario (2) alimenta tablero (3). Tablero individual (3.2) alimenta tickets (4.2). "
        "Tickets en terreno (4.4) alimentan el primer reporte (5.4). Privacidad y despliegue (6) "
        "corren en paralelo, pero 6.4 espera la carga 2.3. Capacitación (7) espera software usable. "
        "Esa lógica —no un plazo “que se nos ocurrió”— es la ruta crítica "
        "1.1 → 2.1 → 2.2 → 2.3 → 3.1 → 3.2 → 4.2 → 4.3 → 4.4 → 5.4 → 7.3.",
        first_line=True,
    )

    path = OUT / "Alcance-EDT-y-exclusiones.docx"
    try:
        doc.save(path)
    except PermissionError:
        path = OUT / "Alcance-EDT-y-exclusiones-actualizado.docx"
        doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    info = run()
    informe = build_informe(info)
    print("INFORME", informe)
    a = build_alcance(info)
    print("ALCANCE", a)
    print("XML", info["xml"])
    print("HTML", info["html"])


if __name__ == "__main__":
    main()
